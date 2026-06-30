#!/usr/bin/env python3
"""Insert BT execution tree table only (flowchart already done)."""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCX_PATH = r"D:\testcase_web\Bao_Cao_Thuc_Tap_WebSec_Test_Behavior_Tree.docx"

def make_run(paragraph, text, font_name="Times New Roman", font_size=13, bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), "Times New Roman")
    if color: run.font.color.rgb = color
    return run

def shade_cell(cell, color):
    tc = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tc.append(shd)

def set_cell(cell, text, bold=False, font_size=9, color=None, align=None):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if align: p.alignment = align
    pf = p.paragraph_format; pf.space_before = Pt(1); pf.space_after = Pt(1)
    make_run(p, text, bold=bold, font_size=font_size, color=color)

def insert_heading_before(doc, ref_idx, text):
    """Insert heading paragraph before ref paragraph, returns the new paragraph."""
    ref_p = doc.paragraphs[ref_idx]
    new_elem = copy.deepcopy(ref_p._element)
    for r in new_elem.findall(qn('w:r')): new_elem.remove(r)
    for rPr in new_elem.findall(qn('w:rPr')): new_elem.remove(rPr)
    ref_p._element.addprevious(new_elem)
    from docx.text.paragraph import Paragraph
    wrapper = Paragraph(new_elem, ref_p._element.getparent())
    wrapper.clear()
    pf = wrapper.paragraph_format; pf.space_before = Pt(8); pf.space_after = Pt(4)
    make_run(wrapper, text, bold=True, font_size=11)
    return wrapper

def main():
    doc = Document(DOCX_PATH)
    print(f"Before: {len(doc.paragraphs)} paras, {len(doc.tables)} tables")

    # Find 3.1.2 heading (with diacritics)
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("3.1.2. Mã nguồn"):
            target_idx = i
            break

    if target_idx is None:
        print("Could not find 3.1.2 heading!")
        return

    print(f"Inserting BT tree after para {target_idx}")

    # Insert heading first
    h = insert_heading_before(doc, target_idx + 1,
        "Hình 3.2: Cấu trúc Behavior Tree kiểm thử đầy đủ (full_scan)")

    # Create table
    rows = [
        ("ROOT: SequenceNode (full_scan)", "Thực thi tuần tự các module", "E2EFDA"),
        ("  +-- ModuleAdapter (headers)", "Kiểm tra 8 Security Headers", "E2EFDA"),
        ("  +-- Retry (auth, attempts=3)", "Thử lại 3 lần nếu auth thất bại", "DAEEF3"),
        ("  |   +-- ModuleAdapter (auth)", "Form login, bypass, rate-limit", "DAEEF3"),
        ("  +-- Selector (fallback)", "Thử disclosure, auth backup", "FDE9D9"),
        ("  |   +-- ModuleAdapter (disclosure)", "Info disclosure check", "FDE9D9"),
        ("  |   +-- ModuleAdapter (auth)", "Phương án dự phòng", "FDE9D9"),
        ("  +-- Parallel (injections)", "Chạy song song 3 injection tests", "E4DFEC"),
        ("  |   +-- SQLi / XSS / CMD inj", "", "E4DFEC"),
        ("  +-- ModuleAdapter (authz)", "Forced browsing + IDOR test", "FCE4D6"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(6.5)

    for i, (label, content, color) in enumerate(rows):
        shade_cell(table.rows[i].cells[0], color)
        shade_cell(table.rows[i].cells[1], color)
        is_bold = not label.startswith("  |   ") and label != "  +-- Parallel (injections)"
        set_cell(table.rows[i].cells[0], label, bold=is_bold, font_size=9,
                 color=RGBColor(0x33, 0x33, 0x33))
        if content:
            set_cell(table.rows[i].cells[1], content, font_size=8.5,
                     color=RGBColor(0x55, 0x55, 0x55))

    # Move table to right after heading
    table._element.getparent().remove(table._element)
    h._element.addnext(table._element)

    doc.save(DOCX_PATH)
    doc2 = Document(DOCX_PATH)
    print(f"After: {len(doc2.paragraphs)} paras, {len(doc2.tables)} tables")
    print("Done.")

if __name__ == "__main__":
    main()
