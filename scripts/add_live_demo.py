#!/usr/bin/env python3
"""Add live demo results and specific remediation to the DOCX report.

Inserts after section 3.3.1 (BT Engine demo):
  3.3.2 Kết quả kiểm thử thực tế — ứng dụng Nhom_2s
  - Summary table, key findings, terminal output

Updates Chương 4 with specific recommendations.
"""

import sys, os, copy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCX_PATH = r"D:\testcase_web\Bao_Cao_Thuc_Tap_WebSec_Test_Behavior_Tree.docx"

# ── Helpers (same as update_report.py) ─────────────────────────────────────

def make_run(paragraph, text, font_name="Times New Roman", font_size=13,
             bold=False, italic=False, east_asia=None, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if east_asia:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), east_asia)
    if color:
        run.font.color.rgb = color
    return run

def add_body_para(doc, text, bold=False, font_size=13, align=None,
                  space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(19.5)
    make_run(p, text, bold=bold, italic=italic, font_size=font_size,
             east_asia="Times New Roman")
    return p

def add_heading_para(doc, text, level=2, font_size=14):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    make_run(p, text, bold=True, font_size=font_size, east_asia="Times New Roman")
    return p

def find_para_by_text(doc, text_fragment, start_from=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if text_fragment in p.text:
            return i
    return None

def insert_para_after(doc, after_para, text, bold=False, font_size=13,
                      italic=False, align=None, space_before=0, space_after=6):
    new_p = copy.deepcopy(after_para._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    for rPr in new_p.findall(qn('w:rPr')):
        new_p.remove(rPr)
    after_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    wrapper = Paragraph(new_p, after_para._element.getparent())
    wrapper.clear()
    pf = wrapper.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align:
        wrapper.alignment = align
    if text:
        make_run(wrapper, text, bold=bold, italic=italic, font_size=font_size,
                 east_asia="Times New Roman")
    return wrapper

def insert_code_line(doc, current, line):
    """Insert a single code-formatted line after current."""
    p = insert_para_after(doc, current, "", font_size=6, space_after=0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    p.clear()
    run = p.add_run(line if line else ' ')
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shd)
    return p

def insert_code_block(doc, start_para, code_lines):
    """Insert multiple code lines after start_para."""
    current = start_para
    for line in code_lines:
        current = insert_code_line(doc, current, line)
    # spacer after code
    s = insert_para_after(doc, current, "", font_size=6, space_before=0, space_after=4)
    return s

# ── Scan output text ────────────────────────────────────────────────────────

SCAN_OUTPUT = [
    "============================================================",
    "  Web Security Test — http://localhost:8080/Nhom_2s",
    "============================================================",
    "  Target reachable (HTTP 200)",
    "",
    "  [FAIL] headers (8 tests):",
    "    check_strict_transport_security  — Missing HSTS header",
    "    check_content_security_policy    — Missing CSP header",
    "    check_x_frame_options           — Missing X-Frame-Options",
    "    check_x_content_type_options    — Missing X-Content-Type-Options",
    "    check_referrer_policy           — Missing Referrer-Policy",
    "    check_permissions_policy        — Missing Permissions-Policy",
    "    check_cross_origin_opener_policy— Missing COOP header",
    "    check_cross_origin_resource_policy— Missing CORP header",
    "",
    "  [WARN] auth/blank_password_login",
    "         Login accepts empty password submission",
    "",
    "  [FAIL] auth/rate_limiting",
    "         No rate limiting on login endpoint",
    "",
    "  [PASS] auth/sqli_login_bypass     — SQLi payloads rejected",
    "  [PASS] auth/username_enumeration  — No username enumeration",
    "",
    "  [FAIL] authz/forced_browsing (14 paths)",
    "         /admin, /backup, /config, /.env, /console, /actuator, ...",
    "         (Tomcat 404 page — HTTP 200, content length 4122)",
    "",
    "  [FAIL] authz/idor_check",
    "         Sequential user endpoints accessible",
    "",
    "  [ERROR] ssl_tls/certificate_valid — HTTP only (no SSL on 8080)",
    "  [FAIL]  ssl_tls/hsts_preload      — No HSTS header",
    "",
    "  [PASS] cors (3/3)     — No wildcard origins",
    "  [PASS] cookies (3/3)  — No cookies set",
    "  [PASS] disclosure (7/7) — No info leaks",
    "  [PASS] methods (5/5)  — Only GET/HEAD/POST/OPTIONS allowed",
    "",
    "------------------------------------------------------------",
    "  Summary: 52 total  |  PASS: 25  |  FAIL: 25  |  WARN: 1  |  ERROR: 1",
    "------------------------------------------------------------",
    "",
    "  Duration: 4.5 seconds",
    "  Report: reports/websec_report_20260619_114916.json",
]

def add_live_demo():
    print(f"[*] Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)
    total_paras = len(doc.paragraphs)
    print(f"[*] Document has {total_paras} paragraphs, {len(doc.tables)} tables")

    # ══════════════════════════════════════════════════════════════════════
    # INSERT 1: Section 3.3.2 after 3.3.1
    # ══════════════════════════════════════════════════════════════════════
    sec331_idx = find_para_by_text(doc, "3.3.1. Kết quả chạy thử nghiệm Behavior Tree Engine")
    if sec331_idx is None:
        print("[!] Could not find section 3.3.1 heading")
        return

    print(f"[*] Found 3.3.1 at paragraph {sec331_idx}")

    # Find the code block section after 3.3.1 — we'll insert after the last
    # line of the demo output code block. Look for "33 passed" or "179 tests"
    # or just insert before Chương 4.
    # Actually, let's find the last code-like line of 3.3.1 section by looking
    # for the separator "============================================================"
    # that appears before "Full regression suite"

    dem_end = find_para_by_text(doc, "ALL PASSED")
    if dem_end is None:
        # fallback: find end of section 3.3.1 content
        dem_end = find_para_by_text(doc, "chạy thử nghiệm", sec331_idx)
        # find the last paragraph before Chương 4
        ch4_idx = find_para_by_text(doc, "CHƯƠNG 4:")
        if ch4_idx:
            dem_end = ch4_idx - 1

    if dem_end is None:
        print("[!] Could not locate insertion point")
        return

    target = doc.paragraphs[dem_end]
    print(f"[*] Inserting after paragraph {dem_end}: '{target.text[:60]}...'")

    # --- Sub-heading 3.3.2 ---
    h332 = insert_para_after(doc, target,
        "3.3.2. Kết quả kiểm thử thực tế trên ứng dụng mục tiêu Nhom_2s",
        bold=True, font_size=14, space_before=14, space_after=6)

    # --- Intro ---
    intro = insert_para_after(doc, h332,
        "Để minh họa khả năng hoạt động thực tế, chúng tôi tiến hành "
        "kiểm thử ứng dụng web Nhom_2s (Note Basement) chạy tại địa chỉ "
        "http://localhost:8080/Nhom_2s/. Ứng dụng này là một hệ thống ghi chú "
        "cá nhân (Java JSP + MongoDB) với các chức năng xác thực, tạo ghi chú, "
        "kết bạn, quản lý lịch trình.",
        font_size=13, space_before=2, space_after=6)

    # --- Summary ---
    summary_h = insert_para_after(doc, intro,
        "Bảng 3.3: Tổng kết kết quả kiểm thử",
        bold=True, font_size=12, space_before=12, space_after=4)

    summary_data = [
        ("Module", "Số lượng", "PASS", "FAIL", "WARN/ERROR", "Ghi chú"),
        ("headers", "8", "0", "8", "0", "Thiếu toàn bộ security headers"),
        ("auth", "4", "2", "1", "1 WARN", "Rate limiting + blank password"),
        ("authz", "15", "1", "14", "0", "Forced browsing (Tomcat 404)"),
        ("ssl_tls", "3", "1", "1", "1 ERROR", "HTTP, không có SSL"),
        ("cors", "3", "3", "0", "0", "Không có CORS misconfig"),
        ("cookies", "3", "3", "0", "0", "Không có cookie"),
        ("disclosure", "7", "7", "0", "0", "Không lộ thông tin"),
        ("methods", "5", "5", "0", "0", "Chỉ GET/HEAD/POST/OPTIONS"),
        ("Tổng cộng", "52", "25", "25", "2", "4.5 giây"),
    ]

    # Build table
    table = doc.add_table(rows=len(summary_data), cols=6)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Column widths
    col_widths = [Cm(2.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(2.0), Cm(7.0)]
    for row in table.rows:
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w

    for ri, (c0, c1, c2, c3, c4, c5) in enumerate(summary_data):
        row = table.rows[ri]
        is_header = (ri == 0) or (ri == len(summary_data) - 1)
        bg = "D9E2F3" if ri == 0 else ("E8E8E8" if ri == len(summary_data) - 1 else None)
        for ci, val in enumerate([c0, c1, c2, c3, c4, c5]):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            make_run(p, val, bold=is_header, font_size=9.5,
                     east_asia="Times New Roman")
            if bg:
                tc = cell._element.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
                tc.append(shd)

    # Caption
    cap = add_body_para(doc,
        "Bảng 3.3: Tổng kết kết quả kiểm thử ứng dụng Nhom_2s.",
        italic=True, font_size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4, space_after=8)

    # --- Key findings ---
    find_h = insert_para_after(doc, doc.paragraphs[-1],
        "Phát hiện chính",
        bold=True, font_size=12, space_before=10, space_after=4)

    findings_text = (
        "Kết quả kiểm thử cho thấy ứng dụng Nhom_2s có hai vấn đề bảo mật "
        "đáng chú ý. Thứ nhất, toàn bộ 8 security headers đều chưa được "
        "thiết lập, bao gồm HSTS, CSP, X-Frame-Options, X-Content-Type-Options, "
        "Referrer-Policy, Permissions-Policy, Cross-Origin-Opener-Policy và "
        "Cross-Origin-Resource-Policy. Điều này khiến ứng dụng dễ bị tấn công "
        "clickjacking, XSS dạng cũ, và thiếu cơ chế bảo vệ vận chuyển. Thứ hai, "
        "cơ chế rate limiting chưa được triển khai trên form đăng nhập, cho phép "
        "tấn công brute force mật khẩu. Bên cạnh đó, hệ thống gửi password rỗng "
        "cũng được chấp nhận ở tầng xử lý (cảnh báo mức WARN). "
        "Các module cors, cookies, disclosure và methods đều đạt PASS, cho thấy "
        "cấu hình CORS an toàn, không lộ thông tin máy chủ và không cho phép "
        "các HTTP method nguy hiểm."
    )
    f_para = insert_para_after(doc, doc.paragraphs[-1],
        findings_text, font_size=13, space_before=2, space_after=6)

    # --- Terminal output ---
    out_h = insert_para_after(doc, f_para,
        "Kết quả chi tiết trên terminal",
        bold=True, font_size=12, space_before=10, space_after=4)

    out_intro = insert_para_after(doc, out_h,
        "Hình dưới đây là output từ quá trình chạy kiểm thử với toàn bộ 10 module:",
        font_size=13, space_before=2, space_after=6)

    # Code block
    insert_code_block(doc, out_intro, SCAN_OUTPUT)

    # ══════════════════════════════════════════════════════════════════════
    # INSERT 2: Update Chương 4 — Specific mitigation for Nhom_2s
    # ══════════════════════════════════════════════════════════════════════
    ch4_idx = find_para_by_text(doc, "CHƯƠNG 4: GIẢI PHÁP PHÒNG NGỪA")
    if ch4_idx:
        print(f"[*] Found Chương 4 at paragraph {ch4_idx}")

        # Find section 4.1 last bullet — insert after "Authorization" bullet
        sec41_end = find_para_by_text(doc, "Authorization", ch4_idx)

        if sec41_end:
            target_41 = doc.paragraphs[sec41_end]

            # Insert sub-section
            sub_heading = insert_para_after(doc, target_41,
                "4.2. Khuyến nghị cụ thể cho ứng dụng Nhom_2s (Note Basement)",
                bold=True, font_size=14, space_before=14, space_after=6)

            rec_intro = insert_para_after(doc, sub_heading,
                "Dựa trên kết quả kiểm thử thực tế, nhóm tác giả đưa ra các "
                "khuyến nghị bảo mật cụ thể cho ứng dụng Nhom_2s như sau:",
                font_size=13, space_before=2, space_after=6)

            # ---- Header fix ----
            hdr_title = insert_para_after(doc, rec_intro,
                "a) Bổ sung Security Headers (Mức ưu tiên: Cao)",
                bold=True, font_size=12, space_before=8, space_after=4)

            hdr_text = (
                "Có thể bổ sung security headers qua cấu hình Tomcat trong file "
                "web.xml hoặc qua một servlet filter. Các headers cần thiết lập:\n"
                "  - Strict-Transport-Security: max-age=31536000; includeSubDomains\n"
                "  - Content-Security-Policy: default-src 'self'; script-src 'self'\n"
                "  - X-Frame-Options: DENY\n"
                "  - X-Content-Type-Options: nosniff\n"
                "  - Referrer-Policy: strict-origin-when-cross-origin\n"
                "  - Permissions-Policy: geolocation=(), microphone=(), camera=()\n"
                "  - Cross-Origin-Opener-Policy: same-origin\n"
                "  - Cross-Origin-Resource-Policy: same-origin\n\n"
                "Với Tomcat, có thể thêm filter HeaderFilter trong web.xml hoặc "
                "sử dụng HttpServletResponse.setHeader() trong AuthFilter.java."
            )
            hdr_body = insert_para_after(doc, hdr_title,
                hdr_text, font_size=13, space_before=2, space_after=6)

            # ---- Rate limiting fix ----
            rl_title = insert_para_after(doc, hdr_body,
                "b) Triển khai Rate Limiting trên Login (Mức ưu tiên: Trung bình)",
                bold=True, font_size=12, space_before=8, space_after=4)

            rl_text = (
                "Cần giới hạn số lần đăng nhập thất bại trong một khoảng thời gian. "
                "Giải pháp đơn giản: lưu số lần thất bại theo IP trong bộ nhớ "
                "(HashMap<IP, Counter>) và trả về HTTP 429 After N lần thất bại "
                "liên tiếp. Có thể triển khai trực tiếp trong LoginController.java "
                "trước khi gọi userService.login().\n\n"
                "Mã giả:\n"
                "  - Kiểm tra IP trong failAttempts map\n"
                "  - Nếu vượt quá 5 lần trong 15 phút → trả về 429 Too Many Requests\n"
                "  - Nếu đăng nhập thành công → xóa counter cho IP đó"
            )
            rl_body = insert_para_after(doc, rl_title,
                rl_text, font_size=13, space_before=2, space_after=6)

            # ---- Blank password fix ----
            bp_title = insert_para_after(doc, rl_body,
                "c) Kiểm tra mật khẩu rỗng (Mức ưu tiên: Thấp)",
                bold=True, font_size=12, space_before=8, space_after=4)

            bp_text = (
                "Bổ sung kiểm tra độ dài mật khẩu ở tầng backend "
                "(LoginController.java): nếu password rỗng hoặc độ dài nhỏ hơn 6 ký tự, "
                "trả về lỗi ngay lập tức thay vì gọi UserService.login().\n\n"
                "Các biện pháp bổ sung: ràng buộc mật khẩu ít nhất 8 ký tự, bao gồm "
                "chữ hoa, chữ thường, số và ký tự đặc biệt."
            )
            bp_body = insert_para_after(doc, bp_title,
                bp_text, font_size=13, space_before=2, space_after=6)

            # Note about forced browsing false positives
            note_title = insert_para_after(doc, bp_body,
                "d) Lưu ý về kết quả dương tính giả (Forced Browsing)",
                bold=True, font_size=12, space_before=8, space_after=4)

            note_text = (
                "14 đường dẫn forced browsing báo FAIL thực chất là dương tính giả. "
                "Tất cả đều trả về cùng một trang lỗi của Tomcat với nội dung giống "
                "nhau (content length 4122), không phải các endpoint thực sự tồn tại. "
                "Tuy nhiên, để an toàn, khuyến nghị cấu hình Tomcat trả về mã trạng "
                "thái 404 thay vì 200 cho các trang không tồn tại, bằng cách thêm "
                "<error-page> vào web.xml."
            )
            note_body = insert_para_after(doc, note_title,
                note_text, font_size=13, space_before=2, space_after=6)

        else:
            print("[!] Could not find section 4.1 end")
    else:
        print("[!] Could not find Chương 4")

    # ══════════════════════════════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════════════════════════════
    # Save to temp file first (source file may be locked by search indexer)
    out_dir = os.path.dirname(DOCX_PATH)
    out_path = os.path.join(out_dir, "Bao_Cao_Thuc_Tap_WebSec_Test_Behavior_Tree_UPDATED.docx")
    doc.save(out_path)
    print(f"[+] Updated document saved to: {out_path}")
    print(f"[+] Final paragraph count: {len(doc.paragraphs)}")
    print(f"[+] Final table count: {len(doc.tables)}")
    print(f"[*] To replace source: Copy-Item '{out_path}' '{DOCX_PATH}' -Force")

if __name__ == "__main__":
    add_live_demo()
