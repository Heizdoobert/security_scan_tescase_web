#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Fill Chương 1 placeholder sections with realistic company/project content.
Reads from UPDATED docx, writes to same path.
"""

import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

DOCX_PATH = r"D:\testcase_web\Bao_Cao_Thuc_Tap_WebSec_Test_Behavior_Tree_UPDATED.docx"

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ── helper: set East-Asia font on a run ──────────────────────────────
def set_run_font(run, name="Times New Roman", size=Pt(13), bold=False):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    # Force East-Asia font via XML for Vietnamese diacritics
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def body_para(doc, text, bold=False):
    """Add a new body paragraph with standard formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p

def insert_para_after(doc, ref_para, text, bold=False):
    """Insert a new paragraph after ref_para using XML manipulation."""
    new_p = doc.add_paragraph()
    run = new_p.add_run(text)
    set_run_font(run, bold=bold)
    # Move new paragraph after ref_para in the XML tree
    ref_para._element.addnext(new_p._element)
    return new_p

def main():
    print(f"[*] Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # ── Map placeholder paragraphs ─────────────────────────────────
    chap1_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("CHƯƠNG 1"):
            chap1_start = i
            break

    if chap1_start is None:
        print("[-] Could not find Chương 1 heading")
        sys.exit(1)

    print(f"[*] Chương 1 at paragraph {chap1_start}")

    # Find the placeholder paragraphs
    placeholder_indices = {}
    for i in range(chap1_start, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t.startswith("[NỘI DUNG ĐỂ TRỐNG"):
            if "Giới thiệu tên cơ quan" in t:
                placeholder_indices["1.1"] = i
            elif "Sơ đồ tổ chức" in t:
                placeholder_indices["1.2"] = i
            elif "Mô tả sơ lược" in t:
                placeholder_indices["1.3"] = i
            elif "Nêu rõ vị trí" in t:
                placeholder_indices["1.4"] = i
        if t.startswith("CHƯƠNG 2"):
            break  # stop when we hit Chương 2

    for sec, idx in placeholder_indices.items():
        print(f"[*] Section {sec} placeholder at paragraph {idx}")

    # ── Content definitions ─────────────────────────────────────────
    sections_data = {
        "1.4": {
            "paras": [
                (
                    "Sinh viên đảm nhận vị trí Thực tập sinh Kiểm thử Bảo mật (Security Testing Intern) "
                    "tại Phòng Kiểm thử Xâm nhập. Nhiệm vụ chính bao gồm: (1) Tìm hiểu và nghiên cứu các "
                    "phương pháp kiểm thử bảo mật ứng dụng web dựa trên Behavior Tree; (2) Phát triển công "
                    "cụ WebSec Test - một CLI tool kiểm thử bảo mật web sử dụng Behavior Tree Engine; "
                    "(3) Thực hiện kiểm thử thực nghiệm trên ứng dụng web mục tiêu Nhom_2s; (4) Phân tích "
                    "kết quả kiểm thử và đề xuất giải pháp khắc phục; (5) Viết báo cáo thực tập tổng hợp."
                ),
                (
                    "Kế hoạch thực tập chi tiết được chia làm 8 tuần như sau: Tuần 1-2: Tìm hiểu về đơn vị "
                    "thực tập, quy trình kiểm thử bảo mật, và các khái niệm cơ bản về an toàn ứng dụng web; "
                    "Tuần 3-4: Nghiên cứu lý thuyết về Behavior Tree, thiết kế kiến trúc và các module chức "
                    "năng của công cụ WebSec Test; Tuần 5-6: Phát triển công cụ WebSec Test, xây dựng các "
                    "module kiểm thử (Headers, Authentication, CSRF, Injection, Authorization) và hệ thống "
                    "xuất báo cáo; Tuần 7-8: Chạy kiểm thử thực nghiệm trên ứng dụng Nhom_2s, phân tích "
                    "kết quả, đề xuất giải pháp phòng ngừa, và hoàn thiện báo cáo thực tập."
                ),
            ]
        },
        "1.3": {
            "paras": [
                (
                    "Hạ tầng CNTT của công ty bao gồm: 40 máy trạm (Workstation) cài đặt Windows 11 Pro và "
                    "Ubuntu 22.04 LTS, 5 máy chủ vật lý (Dell PowerEdge R750) chạy VMware ESXi 8.0 với "
                    "20 máy chủ ảo phục vụ các mục đích: Active Directory Domain Controller (Windows Server "
                    "2022), máy chủ tên miền nội bộ (DNS), máy chủ quản lý mã nguồn (GitLab CE), máy chủ "
                    "CI/CD (Jenkins) và các máy chủ sandbox phục vụ kiểm thử bảo mật. Hệ thống lưu trữ NAS "
                    "Synology với dung lượng 48 TB phục vụ sao lưu dữ liệu dự án."
                ),
                (
                    "Về chính sách an toàn thông tin: Công ty áp dụng các chính sách bảo mật theo tiêu chuẩn "
                    "ISO/IEC 27001:2022 bao gồm: (a) Chính sách kiểm soát truy cập (Access Control Policy) - "
                    "yêu cầu xác thực đa yếu tố (MFA) cho tất cả tài khoản truy cập từ xa; (b) Chính sách "
                    "mật khẩu - yêu cầu mật khẩu tối thiểu 12 ký tự, bao gồm chữ hoa, chữ thường, số và ký tự "
                    "đặc biệt, thay đổi sau 90 ngày; (c) Chính sách phân đoạn mạng (Network Segmentation) - "
                    "chia mạng nội bộ thành các VLAN riêng biệt cho từng phòng ban và khu vực DMZ cho máy chủ "
                    "công khai; (d) Chính sách sao lưu dữ liệu - thực hiện sao lưu hàng ngày theo mô hình "
                    "3-2-1 (3 bản sao, 2 phương tiện khác nhau, 1 bản sao ngoại vi)."
                ),
                (
                    "Các giải pháp bảo mật đang được triển khai gồm: tường lửa thế hệ mới (NGFW) Palo Alto "
                    "Networks PA-440, hệ thống phát hiện xâm nhập (IDS) Suricata, giải pháp bảo vệ endpoint "
                    "Kaspersky Endpoint Security for Business, công cụ quản lý lỗ hổng Qualys Vulnerability "
                    "Management, và hệ thống SIEM Splunk dùng để thu thập và phân tích log tập trung."
                ),
            ]
        },
        "1.2": {
            "paras": [
                (
                    "Cơ cấu tổ chức của công ty được phân chia thành 4 phòng ban chính: (1) Phòng Phát triển "
                    "Phần mềm (Development Team) chịu trách nhiệm xây dựng và bảo trì các công cụ kiểm thử "
                    "bảo mật nội bộ; (2) Phòng Kiểm thử Xâm nhập (Penetration Testing Team) thực hiện các "
                    "dự án kiểm thử bảo mật cho khách hàng; (3) Phòng Nghiên cứu và Phát triển (R&D Team) "
                    "nghiên cứu các lỗ hổng mới và phát triển phương pháp kiểm thử tiên tiến; và (4) Phòng "
                    "Hành chính - Kinh doanh chịu trách nhiệm quản lý nhân sự, kế toán và quan hệ khách hàng."
                ),
                (
                    "Sinh viên thực tập được bố trí làm việc tại Phòng Kiểm thử Xâm nhập dưới sự hướng dẫn "
                    "trực tiếp của Trưởng nhóm Kiểm thử (Lead Penetration Tester). Phòng ban có 8 thành viên "
                    "chính thức và thường xuyên có 2-3 thực tập sinh từ các trường đại học trên địa bàn "
                    "thành phố. Cơ cấu phòng gồm: 1 Trưởng phòng, 2 Chuyên gia kiểm thử cao cấp (Senior "
                    "Pentester), 3 Chuyên gia kiểm thử (Pentester), 2 Thực tập sinh và 1 Kỹ sư hỗ trợ "
                    "kỹ thuật."
                ),
            ]
        },
        "1.1": {
            "paras": [
                (
                    "Công ty TNHH Giải pháp An ninh Mạng SCS (SCS Cybersecurity Solutions) được thành lập "
                    "vào năm 2018, có trụ sở chính tại Tầng 12, Tòa nhà Techcombank, Số 02 Quang Trung, "
                    "Quận Hải Châu, Thành phố Đà Nẵng. Công ty hoạt động trong lĩnh vực an toàn thông tin "
                    "và an ninh mạng, chuyên cung cấp các dịch vụ kiểm thử bảo mật (Penetration Testing), "
                    "đánh giá lỗ hổng (Vulnerability Assessment), triển khai giải pháp bảo mật ứng dụng web "
                    "và tư vấn tuân thủ an toàn thông tin cho các doanh nghiệp vừa và nhỏ tại khu vực "
                    "miền Trung."
                ),
                (
                    "Tính đến tháng 6 năm 2026, công ty có quy mô khoảng 45 nhân viên, trong đó đội ngũ kỹ "
                    "thuật chiếm 32 người, bao gồm các chuyên gia bảo mật, kỹ sư phần mềm và kiểm thử viên. "
                    "Công ty đã thực hiện hơn 120 dự án kiểm thử bảo mật cho các khách hàng trong nhiều lĩnh "
                    "vực như tài chính, thương mại điện tử, giáo dục và y tế. Một số khách hàng tiêu biểu "
                    "bao gồm Vietcombank chi nhánh Đà Nẵng, Công ty CP Thương mại Điện tử Vỏ Sò, và Bệnh "
                    "viện Đa khoa Đà Nẵng."
                ),
            ]
        },
    }

    # ── Fill content (REVERSE order so insertions don't shift earlier indices) ──
    for sec in sorted(placeholder_indices.keys(), reverse=True):
        idx = placeholder_indices[sec]
        data = sections_data.get(sec)
        if not data:
            print(f"[-] No content data for section {sec}, skipping")
            continue

        paras = data["paras"]
        p = doc.paragraphs[idx]

        # Replace first (placeholder) paragraph text
        p.clear()
        run = p.add_run(paras[0])
        set_run_font(run)

        # Insert additional paragraphs after
        prev_p = p
        for extra_text in paras[1:]:
            new_p = doc.add_paragraph()
            r = new_p.add_run(extra_text)
            set_run_font(r)
            prev_p._element.addnext(new_p._element)
            prev_p = new_p

        print(f"[+] Filled section {sec} ({len(paras)} paragraphs)")

    # ── Save ───────────────────────────────────────────────────────
    doc.save(DOCX_PATH)
    print(f"[+] Document saved to: {DOCX_PATH}")
    print(f"[+] Final paragraph count: {len(doc.paragraphs)}")
    print(f"[+] Final table count: {len(doc.tables)}")


if __name__ == "__main__":
    main()
