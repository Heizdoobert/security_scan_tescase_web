#!/usr/bin/env python3
"""Insert architecture flowchart and BT execution tables into the report."""
import copy, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCX_PATH = r"D:\testcase_web\Bao_Cao_Thuc_Tap_WebSec_Test_Behavior_Tree.docx"

def make_run(paragraph, text, font_name="Times New Roman", font_size=13,
             bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), "Times New Roman")
    if color:
        run.font.color.rgb = color
    return run

def shade_cell(cell, color):
    tc = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tc.append(shd)

def set_cell_text(cell, text, bold=False, font_size=9.5, color=None, align=None):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    make_run(p, text, bold=bold, font_size=font_size, color=color)

def insert_para_after_ref(doc, ref_idx, text, bold=False, font_size=13,
                          italic=False, align=None, space_before=0, space_after=6):
    """Insert paragraph after ref paragraph (reference by index)."""
    ref_p = doc.paragraphs[ref_idx]
    new_elem = copy.deepcopy(ref_p._element)
    for r in new_elem.findall(qn('w:r')): new_elem.remove(r)
    for rPr in new_elem.findall(qn('w:rPr')): new_elem.remove(rPr)
    ref_p._element.addnext(new_elem)
    from docx.text.paragraph import Paragraph
    wrapper = Paragraph(new_elem, ref_p._element.getparent())
    wrapper.clear()
    pf = wrapper.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align: wrapper.alignment = align
    if text:
        make_run(wrapper, text, bold=bold, italic=italic, font_size=font_size)
    return wrapper

def add_flowchart(doc, insert_after_idx):
    """Insert architecture flowchart table."""
    insert_para_after_ref(doc, insert_after_idx, "", font_size=6, space_after=2)
    # heading
    p = insert_para_after_ref(doc, insert_after_idx,
        "Hình 3.1: Kiến trúc 3 tầng của WebSec Test", bold=True, font_size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    rows_data = [
        ("Tầng 1: CLI Entry Point", "argparse -> main.py\nPhân tích tham số, điều phối luồng kiểm thử", "D4E8F0"),
        ("", "  |  ", ""),
        ("Tầng 2: Core Engine", "SessionClient + Behavior Tree Engine\nBlackboard o Sequence o Selector o Parallel o Decorators", "D5F0D5"),
        ("", "  |  ", ""),
        ("Tầng 3: Modules & Results", "10 Module kiểm thử -> ResultCollector -> Reporter\nheaders, auth, csrf, sql_injection, authz, ssl_tls...", "FCE4D6"),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, content, color) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11)
        if label:
            shade_cell(row.cells[0], color)
            shade_cell(row.cells[1], color)
            set_cell_text(row.cells[0], label, bold=True, font_size=10)
            set_cell_text(row.cells[1], content, font_size=10)
        else:
            row.cells[0].merge(row.cells[1])
            set_cell_text(row.cells[0], content, bold=True, font_size=14,
                         color=RGBColor(0x44, 0x72, 0xC4), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Move table after the heading
    # (doc.add_table appends at end; we need to move it)
    table._element.getparent().remove(table._element)
    p._element.addnext(table._element)
    return table

def add_bt_tree(doc, insert_after_idx):
    """Insert BT execution tree table."""
    insert_para_after_ref(doc, insert_after_idx, "", font_size=6, space_after=2)
    p = insert_para_after_ref(doc, insert_after_idx,
        "Hình 3.2: Cấu trúc Behavior Tree kiểm thử đầy đủ (full_scan)", bold=True,
        font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    rows_data = [
        ("ROOT: SequenceNode (full_scan)", "Thực thi tuần tự các module", "E2EFDA"),
        ("  +-- ModuleAdapter (headers)", "Kiểm tra 8 Security Headers", "E2EFDA"),
        ("  +-- Retry (auth, attempts=3)", "Thử lại 3 lan auth", "DAEEF3"),
        ("  |   +-- ModuleAdapter (auth)", "Form login, bypass, rate-limit", "DAEEF3"),
        ("  +-- Selector (fallback)", "Thu: disclosure > auth backup", "FDE9D9"),
        ("  |   +-- ModuleAdapter (disclosure)", "Info disclosure check", "FDE9D9"),
        ("  |   +-- ModuleAdapter (auth)", "Phuong an du phong", "FDE9D9"),
        ("  +-- Parallel (injections)", "Song song 3 injection tests", "E4DFEC"),
        ("  |   +-- SQL Injection / XSS / CMD", "", "E4DFEC"),
        ("  +-- ModuleAdapter (authz)", "Forced browsing + IDOR", "FCE4D6"),
    ]
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, content, color) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].width = Cm(9)
        row.cells[1].width = Cm(6.5)
        shade_cell(row.cells[0], color)
        shade_cell(row.cells[1], color)
        is_bold = not label.startswith("  |   ") and not label.startswith("  +-- Parallel")
        c = RGBColor(0x33, 0x33, 0x33)
        set_cell_text(row.cells[0], label, bold=is_bold, font_size=9, color=c)
        if content:
            set_cell_text(row.cells[1], content, font_size=8.5,
                         color=RGBColor(0x55, 0x55, 0x55))

    # Move after heading
    table._element.getparent().remove(table._element)
    p._element.addnext(table._element)
    return table

def main():
    doc = Document(DOCX_PATH)
    total = len(doc.paragraphs)

    # Find insertion points by section content
    # Flowchart goes after "3.1.1. Behavior Tree Engine - Thiet ke chi tiet" heading
    insert_flow = None
    insert_tree = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "3.1.1. Behavior Tree Engine" in t:
            insert_flow = i
        if "3.1.2. Ma nguon trien khai" in t:
            insert_tree = i

    print(f"Insert flow after para {insert_flow} (3.1.1)")
    print(f"Insert tree after para {insert_tree} (3.1.2)")

    if insert_flow:
        add_flowchart(doc, insert_flow)
        print("+ Added architecture flowchart (Hinh 3.1)")

    if insert_tree:
        add_bt_tree(doc, insert_tree)
        print("+ Added BT execution tree (Hinh 3.2)")

    doc.save(DOCX_PATH)
    print(f"Saved: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

if __name__ == "__main__":
    main()
